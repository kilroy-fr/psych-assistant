# app/docx_generator.py
"""DOCX-Generierung fuer Psychotherapie-Berichte."""

import io
import os
import re
import json
from docx import Document
from docx.shared import Pt, Cm, Twips
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.enum.style import WD_STYLE_TYPE

TEMPLATE_FILENAME = "report_template.docx"

# Lade Schema-Definitionen für Überschriften
def load_schema():
    """Lädt das Report-Schema für Überschriftenstruktur."""
    schema_path = os.path.join(
        os.path.dirname(__file__),
        "..",
        "data",
        "guidelines",
        "report_schema_vt_umwandlung.json"
    )
    try:
        with open(schema_path, "r", encoding="utf-8") as f:
            return json.load(f)
    except FileNotFoundError:
        return None

# Erstelle Überschriften-Mapping aus Schema
def build_heading_map(schema):
    """Erstellt ein Mapping von Überschriften zu Hierarchieebenen.

    Returns:
        dict: {überschrift: level} wobei level = 2 (Hauptüberschrift), 3 (Unterüberschrift)

    Unterstützt auch "alternatives" Array im Schema für geschlechtsspezifische Varianten.
    """
    if not schema:
        return {}

    heading_map = {}
    for section in schema.get("sections", []):
        # Hauptüberschrift (Level 2 = h2)
        heading_map[section["title"]] = 2
        # Alternative Hauptüberschriften
        for alt in section.get("alternatives", []):
            heading_map[alt] = 2

        # Unterüberschriften (Level 3 = h3)
        for subsection in section.get("subsections", []):
            heading_map[subsection["title"]] = 3
            # Alternative Unterüberschriften (z.B. männliche Varianten)
            for alt in subsection.get("alternatives", []):
                heading_map[alt] = 3

    return heading_map

# Globales Schema und Heading-Map laden
SCHEMA = load_schema()
HEADING_MAP = build_heading_map(SCHEMA)

SENSITIVE_PATTERNS = [
    (re.compile(r"\[Anonymisiert\]"), "X."),
    (re.compile(r"\b(Frau|Herr)\s+[A-ZÄÖÜ][a-zäöüß-]+"), r"\1 X."),
    (re.compile(r"\b(Frau|Herr)\s+[A-ZÄÖÜ]\."), r"\1 X."),
    (re.compile(r"\b(Name|Vorname|Nachname|Geburtsname)\b\s*[:\-]\s*[^\n]+"), r"\1: X."),
    (re.compile(r"\b(Ort|Wohnort|Geburtsort|Adresse|Straße|Strasse|Stadt|PLZ)\b\s*[:\-]\s*[^\n]+"), r"\1: F."),
    (re.compile(r"\b[A-ZÄÖÜ]\.\s*[A-ZÄÖÜ]\.\b"), "X."),
]


# =============================================================================
# A) DETERMINISTISCHE FORMAT-SÄUBERUNG
# =============================================================================

def clean_output(text):
    """Entfernt Markdown-Artefakte und Listenoptik deterministisch.

    Diese Funktion wird vor sanitize_sensitive_text() ausgeführt und garantiert,
    dass kein Markdown mehr im Output verbleibt, unabhängig vom LLM-Output.

    Regel 1: Entferne Markdown-Überschriften am Zeilenanfang (^#{1,6}\s*)
    Regel 2: Entferne Bullet-Startzeichen am Zeilenanfang (^[-*•]\s+)
    Regel 3: Entferne Markdown-Formatierungen im Fließtext (**text**, *text*)
    Regel 4: Normalisiere Leerzeilen (max. eine aufeinanderfolgende Leerzeile)
    Regel 5: Entferne isolierte Markdown-Artefakte (einzelne #, ##, etc.)

    Args:
        text: Der zu säubernde Text

    Returns:
        str: Gesäuberter Text ohne Markdown und Listenoptik
    """
    if not text:
        return text

    lines = text.splitlines()
    cleaned_lines = []

    for line in lines:
        # Regel 1: Entferne Markdown-Überschriften am Zeilenanfang
        line = re.sub(r"^#{1,6}\s*", "", line)

        # Regel 2: Entferne Bullet-Startzeichen am Zeilenanfang
        line = re.sub(r"^[-*•]\s+", "", line)

        # Regel 3: Entferne Markdown-Formatierungen im Fließtext
        # **text** → text (Bold)
        line = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
        # *text* → text (Italic) - aber nur wenn nicht Teil von **
        line = re.sub(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)", r"\1", line)

        # Regel 5: Entferne isolierte Markdown-Artefakte (nur #-Zeichen)
        if re.match(r"^#+$", line.strip()):
            continue  # Überspringe diese Zeile komplett

        cleaned_lines.append(line)

    # Regel 4: Normalisiere Leerzeilen (max. eine aufeinanderfolgende)
    result_lines = []
    prev_was_empty = False

    for line in cleaned_lines:
        is_empty = not line.strip()

        if is_empty:
            if not prev_was_empty:
                result_lines.append(line)
            prev_was_empty = True
        else:
            result_lines.append(line)
            prev_was_empty = False

    return "\n".join(result_lines)


# =============================================================================
# B) SCHEMA-VALIDATOR + REPAIR
# =============================================================================

def normalize_heading(text):
    """Normalisiert Überschriften für Vergleich (entfernt Nummerierung, Whitespace, Groß-/Kleinschreibung).

    Args:
        text: Der zu normalisierende Text

    Returns:
        str: Normalisierter Text (lowercase, keine Nummerierung, kein extra whitespace, normalisierte Umlaute)
    """
    # Entferne Nummerierung am Anfang (z.B. "2.1 ", "4.2. ", "6.1 ", etc.)
    pattern = r"^(?:\d+\.?\d*\.?\s*)?(.+)$"
    match = re.match(pattern, text.strip())
    if match:
        text = match.group(1).strip()

    # Lowercase
    text = text.lower().strip()

    # Normalisiere Umlaute (ä->ae, ö->oe, ü->ue, ß->ss)
    # Dies macht den Vergleich robust gegen unterschiedliche Umlaut-Schreibweisen
    umlaut_map = {
        'ä': 'ae', 'ö': 'oe', 'ü': 'ue', 'ß': 'ss'
    }
    for umlaut, replacement in umlaut_map.items():
        text = text.replace(umlaut, replacement)

    # Normalisiere Whitespace (mehrfache Leerzeichen -> eins)
    return re.sub(r'\s+', ' ', text)


def build_heading_synonyms():
    """Erstellt ein Mapping von Synonymen/Varianten zu kanonischen Schema-Überschriften.

    Returns:
        dict: {variante_normalized: kanonische_überschrift}
    """
    synonyms = {}

    # Explizite Synonyme für bekannte Varianten
    known_variants = {
        "psychischer befund": "Psychopathologischer Befund",
        "psychopathologie": "Psychopathologischer Befund",
        "befund": "Psychopathologischer Befund",
        "somatischer befund / konsiliarbericht": "Somatischer Befund",
        "konsiliarbericht": "Somatischer Befund",
        "diagnose": "Diagnose nach ICD-10",
        "diagnosen": "Diagnose nach ICD-10",
        "icd-10": "Diagnose nach ICD-10",
        "therapieziele": "Therapieziele (mit der Patientin vereinbart)",
        "behandlungsplan": "Individueller Behandlungsplan",
    }

    # Füge bekannte Varianten hinzu
    for variant, canonical in known_variants.items():
        synonyms[normalize_heading(variant)] = canonical

    # Füge alle Schema-Überschriften als ihre eigenen kanonischen Versionen hinzu
    if SCHEMA:
        for section in SCHEMA.get("sections", []):
            normalized = normalize_heading(section["title"])
            synonyms[normalized] = section["title"]

            for subsection in section.get("subsections", []):
                normalized_sub = normalize_heading(subsection["title"])
                synonyms[normalized_sub] = subsection["title"]

    return synonyms


# Globales Synonym-Mapping
HEADING_SYNONYMS = build_heading_synonyms()


class ValidationResult:
    """Ergebnis der Schema-Validierung."""

    def __init__(self):
        self.is_valid = True
        self.errors = []
        self.warnings = []
        self.missing_sections = []
        self.missing_subsections = {}
        self.markdown_artifacts = []
        self.placeholder_locations = []

    def add_error(self, error):
        self.is_valid = False
        self.errors.append(error)

    def add_warning(self, warning):
        self.warnings.append(warning)


def validate_schema(text):
    """Validiert einen Text gegen das Report-Schema.

    Prüft:
    1. Sind alle 6 Hauptabschnitte vorhanden?
    2. Sind alle Pflicht-Unterüberschriften vorhanden?
    3. Enthält der Text noch Markdown-Artefakte?
    4. Enthält der Text Platzhalter wie [Angabe fehlt]?

    Args:
        text: Der zu validierende Text

    Returns:
        ValidationResult: Validierungsergebnis mit Fehlern und Warnungen
    """
    result = ValidationResult()

    if not text or not SCHEMA:
        result.add_error("Text oder Schema nicht verfügbar")
        return result

    lines = text.splitlines()
    normalized_lines = [normalize_heading(line) for line in lines]

    # Prüfung 1 & 2: Hauptabschnitte und Unterüberschriften
    for section in SCHEMA.get("sections", []):
        section_title = section["title"]
        section_normalized = normalize_heading(section_title)

        # Prüfe Hauptabschnitt
        if not any(section_normalized in norm_line for norm_line in normalized_lines):
            result.add_error(f"Hauptabschnitt fehlt: {section_title}")
            result.missing_sections.append(section_title)

        # Prüfe Unterüberschriften (nur wenn required=true)
        for subsection in section.get("subsections", []):
            if subsection.get("required", False):
                subsection_title = subsection["title"]
                subsection_normalized = normalize_heading(subsection_title)

                if not any(subsection_normalized in norm_line for norm_line in normalized_lines):
                    section_id = section["id"]
                    if section_id not in result.missing_subsections:
                        result.missing_subsections[section_id] = []
                    result.missing_subsections[section_id].append(subsection_title)
                    result.add_error(f"Unterüberschrift fehlt in Abschnitt {section_id}: {subsection_title}")

    # Prüfung 3: Markdown-Artefakte (sollte nach clean_output() keine mehr geben, aber double-check)
    for i, line in enumerate(lines, 1):
        # Markdown-Überschriften
        if re.match(r"^#{1,6}\s+", line):
            result.add_warning(f"Zeile {i}: Markdown-Überschrift gefunden: {line[:50]}")
            result.markdown_artifacts.append((i, line))

        # Bullet-Zeichen am Zeilenanfang
        if re.match(r"^[-*•]\s+", line):
            result.add_warning(f"Zeile {i}: Bullet-Zeichen gefunden: {line[:50]}")
            result.markdown_artifacts.append((i, line))

    # Prüfung 4: Platzhalter
    placeholder_pattern = re.compile(r"\[Angabe fehlt\]")
    for i, line in enumerate(lines, 1):
        if placeholder_pattern.search(line):
            result.add_warning(f"Zeile {i}: Platzhalter gefunden")
            result.placeholder_locations.append(i)

    return result


def repair_schema(text):
    """Repariert einen Text basierend auf Schema-Vorgaben.

    Automatische Korrekturen:
    1. Normalisiert Überschriften (Synonyme → kanonische Namen)
    2. Erzwingt Reihenfolge der Abschnitte gemäß Schema
    3. Fügt fehlende Abschnitte mit [Angabe fehlt] ein (außer Abschnitt 5)

    WICHTIG: Inhaltliche Ergänzungen werden NICHT vorgenommen.
    Fehlender Abschnitt 5 wird NICHT automatisch eingefügt (erfordert LLM-Run).

    Args:
        text: Der zu reparierende Text

    Returns:
        tuple: (reparierter_text, repair_log)
    """
    if not text or not SCHEMA:
        return text, []

    repair_log = []
    lines = text.splitlines()

    # Schritt 1: Normalisiere Überschriften (Synonyme → kanonische Namen)
    normalized_lines = []
    for line in lines:
        cleaned = line.strip()
        if not cleaned:
            normalized_lines.append(line)
            continue

        # Prüfe ob es eine Überschrift ist
        heading_level = detect_heading_level(cleaned)
        if heading_level:
            # Suche kanonische Version
            normalized_heading = normalize_heading(cleaned)
            if normalized_heading in HEADING_SYNONYMS:
                canonical = HEADING_SYNONYMS[normalized_heading]
                if canonical != cleaned:
                    repair_log.append(f"Überschrift normalisiert: '{cleaned}' → '{canonical}'")
                    normalized_lines.append(canonical)
                else:
                    normalized_lines.append(line)
            else:
                normalized_lines.append(line)
        else:
            normalized_lines.append(line)

    # Schritt 2: Extrahiere Abschnitte aus dem Text
    sections_found = {}
    current_section_id = None
    current_content = []

    for line in normalized_lines:
        cleaned = line.strip()

        # Prüfe ob es eine Hauptüberschrift ist (Level 2)
        if detect_heading_level(cleaned) == 2:
            # Speichere vorherigen Abschnitt
            if current_section_id:
                sections_found[current_section_id] = "\n".join(current_content)

            # Finde welchem Schema-Abschnitt diese Überschrift entspricht
            current_section_id = None
            for section in SCHEMA.get("sections", []):
                if normalize_heading(section["title"]) == normalize_heading(cleaned):
                    current_section_id = section["id"]
                    current_content = [line]
                    break

            if not current_section_id:
                # Unbekannte Überschrift - füge zu aktuellem Abschnitt hinzu
                if current_content:
                    current_content.append(line)
        else:
            current_content.append(line)

    # Speichere letzten Abschnitt
    if current_section_id:
        sections_found[current_section_id] = "\n".join(current_content)

    # Schritt 3: Baue Text in korrekter Schema-Reihenfolge neu auf
    reconstructed_lines = []

    for section in SCHEMA.get("sections", []):
        section_id = section["id"]

        if section_id in sections_found:
            # Abschnitt vorhanden - übernehme
            reconstructed_lines.append(sections_found[section_id])
        else:
            # Abschnitt fehlt
            if section_id == "5":
                # Abschnitt 5 (Diagnose) NICHT automatisch einfügen
                repair_log.append(f"WARNUNG: Abschnitt 5 (Diagnose) fehlt - erfordert LLM-Run")
            else:
                # Andere fehlende Abschnitte mit Platzhalter einfügen
                repair_log.append(f"Abschnitt {section_id} fehlt - Platzhalter eingefügt")
                reconstructed_lines.append(section["title"])
                reconstructed_lines.append("[Angabe fehlt]")

        # Füge Leerzeile zwischen Abschnitten hinzu
        if section_id != SCHEMA["sections"][-1]["id"]:
            reconstructed_lines.append("")

    return "\n".join(reconstructed_lines), repair_log


# =============================================================================
# C) HAUPT-POST-PROCESSING-PIPELINE
# =============================================================================

def reorganize_subsections(text):
    """Reorganisiert Unterabschnitte basierend auf Schema-Reihenfolge.

    Diese Funktion arbeitet auf Textebene und reorganisiert Subsections,
    die entweder als Überschriften ODER als Fließtext mit Doppelpunkt formatiert sind.

    Beispiel:
    Input:  "3. Somatischer Befund\naktuelle Medikation: ...\nVorbehandlungen: ...\nBefunde: ..."
    Output: "3. Somatischer Befund\nBefunde: ...\naktuelle Medikation: ...\nVorbehandlungen: ..."

    Args:
        text: Der zu reorganisierende Text

    Returns:
        str: Text mit reorganisierten Unterabschnitten
    """
    if not text or not SCHEMA:
        return text

    lines = text.splitlines()
    result_lines = []

    i = 0
    while i < len(lines):
        line = lines[i]
        cleaned = line.strip()

        # Prüfe ob dies eine Hauptüberschrift ist
        heading_level = detect_heading_level(cleaned) if cleaned else None

        if heading_level == 2:
            # Finde Schema-Definition für diesen Hauptabschnitt
            schema_section = None
            for s in SCHEMA.get("sections", []):
                if normalize_heading(s["title"]) == normalize_heading(cleaned):
                    schema_section = s
                    break

            # Hauptüberschrift zur Ausgabe hinzufügen
            result_lines.append(line)
            i += 1

            if not schema_section or 'subsections' not in schema_section:
                # Kein Schema oder keine Subsections - weitermachen
                continue

            # Sammle alle Subsections und deren Inhalt
            subsection_data = {}  # {normalized_title: [lines]}
            current_subsection = None
            current_subsection_lines = []

            # Sammle Zeilen bis zur nächsten Hauptüberschrift
            while i < len(lines):
                next_line = lines[i]
                next_cleaned = next_line.strip()

                # Stop bei nächster Hauptüberschrift
                if next_cleaned and detect_heading_level(next_cleaned) == 2:
                    break

                # Prüfe ob diese Zeile eine Subsection startet
                matched_subsection = None
                for schema_sub in schema_section['subsections']:
                    sub_title = schema_sub['title']
                    normalized_schema = normalize_heading(sub_title)

                    # Prüfe auf Überschrift (Heading 3)
                    if next_cleaned and detect_heading_level(next_cleaned) == 3:
                        normalized_line = normalize_heading(next_cleaned)
                        if normalized_line == normalized_schema:
                            matched_subsection = normalized_schema
                            break

                    # Prüfe auf Fließtext-Format "Titel: Inhalt..."
                    if ':' in next_cleaned:
                        prefix = next_cleaned.split(':', 1)[0].strip()
                        normalized_prefix = normalize_heading(prefix)
                        if normalized_prefix == normalized_schema:
                            matched_subsection = normalized_schema
                            break

                if matched_subsection:
                    # Speichere vorherige Subsection
                    if current_subsection and current_subsection_lines:
                        subsection_data[current_subsection] = current_subsection_lines

                    # Starte neue Subsection
                    current_subsection = matched_subsection
                    current_subsection_lines = [next_line]
                else:
                    # Füge zur aktuellen Subsection hinzu (oder zu "other" wenn keine aktiv)
                    if current_subsection:
                        current_subsection_lines.append(next_line)
                    else:
                        # Inhalt vor erster Subsection
                        result_lines.append(next_line)

                i += 1

            # Speichere letzte Subsection
            if current_subsection and current_subsection_lines:
                subsection_data[current_subsection] = current_subsection_lines

            # Füge Subsections in Schema-Reihenfolge zur Ausgabe hinzu
            for schema_sub in schema_section['subsections']:
                normalized_schema = normalize_heading(schema_sub['title'])
                if normalized_schema in subsection_data:
                    result_lines.extend(subsection_data[normalized_schema])

        else:
            # Normale Zeile - zur Ausgabe hinzufügen
            result_lines.append(line)
            i += 1

    return '\n'.join(result_lines)


def post_process_text(text, enable_repair=True, enable_validation=True):
    """Führt vollständige Post-Processing-Pipeline aus.

    Pipeline-Schritte:
    1. Deterministische Format-Säuberung (clean_output)
    2. Schema-Reparatur (repair_schema) - optional (VORSICHT: kann Abschnitte reorganisieren!)
    3. Schema-Validierung (validate_schema) - optional
    4. Subsection-Reorganisation (reorganize_subsections) - immer
    5. Nummerierung hinzufügen (add_section_numbering) - immer
    6. Sensible Daten anonymisieren (sanitize_sensitive_text) - immer

    Args:
        text: Der zu verarbeitende Text
        enable_repair: Ob automatische Reparaturen durchgeführt werden sollen (Standard: True)
                       WARNUNG: repair_schema reorganisiert den gesamten Text!
        enable_validation: Ob Validierung durchgeführt werden soll (Standard: True)

    Returns:
        dict: {
            "text": verarbeiteter Text,
            "validation": ValidationResult (oder None),
            "repair_log": Liste der Reparatur-Aktionen,
            "cleaned": bool (ob Format-Säuberung durchgeführt wurde)
        }
    """
    result = {
        "text": text,
        "validation": None,
        "repair_log": [],
        "cleaned": False
    }

    if not text:
        return result

    # Schritt 1: Format-Säuberung (immer durchführen)
    cleaned_text = clean_output(text)
    result["cleaned"] = True
    result["text"] = cleaned_text

    # Schritt 2: Schema-Reparatur (optional, vor Validierung)
    if enable_repair:
        repaired_text, repair_log = repair_schema(cleaned_text)
        result["text"] = repaired_text
        result["repair_log"] = repair_log

    # Schritt 3: Schema-Validierung (optional)
    if enable_validation:
        validation_result = validate_schema(result["text"])
        result["validation"] = validation_result

    # Schritt 4: Subsection-Reorganisation (immer durchführen)
    result["text"] = reorganize_subsections(result["text"])

    # Schritt 5: Nummerierung hinzufügen (immer durchführen)
    result["text"] = add_section_numbering(result["text"])

    # Schritt 6: Sensible Daten anonymisieren (immer durchführen)
    result["text"] = sanitize_sensitive_text(result["text"])

    return result


def add_section_numbering(text):
    """Fügt Nummerierung zu Überschriften hinzu und REORGANISIERT Unterabschnitte basierend auf Schema.

    Hauptabschnitte: 1, 2, 3, 4, 5, 6
    Unterabschnitte: 1.1, 2.1, 2.2, etc.

    KRITISCH: Diese Funktion REORGANISIERT Unterabschnitte in die Schema-konforme Reihenfolge!
    Wenn ein LLM z.B. Unterabschnitte in der Reihenfolge 3.3, 3.2, 3.1 generiert,
    werden sie in 3.1, 3.2, 3.3 umgeordnet.

    Args:
        text: Der Text mit Überschriften

    Returns:
        str: Text mit nummerierten und reorganisierten Überschriften
    """
    if not text or not SCHEMA:
        return text

    lines = text.splitlines()
    result_lines = []

    # Parse den Text in Abschnitte mit Unterabschnitten
    sections = []  # Liste von {type: 'main'/'sub'/'content', ...}
    current_section = None

    i = 0
    while i < len(lines):
        line = lines[i]
        cleaned = line.strip()

        if not cleaned:
            if current_section:
                current_section['lines'].append(line)
            else:
                result_lines.append(line)
            i += 1
            continue

        heading_level = detect_heading_level(cleaned)

        if heading_level == 2:
            # Neue Hauptüberschrift
            if current_section:
                sections.append(current_section)

            current_section = {
                'type': 'main',
                'heading': cleaned,
                'lines': [],
                'subsections': []
            }
            i += 1

        elif heading_level == 3:
            # Unterüberschrift
            if current_section and current_section['type'] == 'main':
                # Sammle den Inhalt dieser Unterüberschrift
                subsection_lines = []
                i += 1  # Skip heading line

                # Sammle alle Zeilen bis zur nächsten Überschrift
                while i < len(lines):
                    next_cleaned = lines[i].strip()
                    if next_cleaned and detect_heading_level(next_cleaned) in [2, 3]:
                        break
                    subsection_lines.append(lines[i])
                    i += 1

                current_section['subsections'].append({
                    'heading': cleaned,
                    'lines': subsection_lines
                })
            else:
                # Kein aktiver Hauptabschnitt - als normalen Text behandeln
                if current_section:
                    current_section['lines'].append(line)
                else:
                    result_lines.append(line)
                i += 1
        else:
            # Normaler Text
            if current_section:
                current_section['lines'].append(line)
            else:
                result_lines.append(line)
            i += 1

    # Letzten Abschnitt hinzufügen
    if current_section:
        sections.append(current_section)

    # Jetzt reorganisieren und nummerieren
    for section in sections:
        if section['type'] != 'main':
            continue

        # Finde Schema-Definition für diesen Hauptabschnitt
        schema_section = None
        for s in SCHEMA.get("sections", []):
            if normalize_heading(s["title"]) == normalize_heading(section['heading']):
                schema_section = s
                break

        if not schema_section:
            # Unbekannter Abschnitt - unverändert ausgeben
            result_lines.append(section['heading'])
            result_lines.extend(section['lines'])
            for subsection in section['subsections']:
                result_lines.append(subsection['heading'])
                result_lines.extend(subsection['lines'])
            continue

        # Nummeriere Hauptüberschrift
        clean_title = re.sub(r"^(?:\d+\.?\d*\.?\s*)?(.+)$", r"\1", section['heading'].strip())
        numbered_title = f"{schema_section['id']}. {clean_title}"
        result_lines.append(numbered_title)

        # Füge Inhalte vor Unterabschnitten hinzu
        result_lines.extend(section['lines'])

        # Reorganisiere Unterabschnitte basierend auf Schema
        if 'subsections' in schema_section and section['subsections']:
            schema_subsections = schema_section['subsections']

            # Erstelle Mapping: normalized_title -> subsection_data
            subsection_map = {}
            for subsection in section['subsections']:
                clean_sub_title = re.sub(r"^(?:\d+\.?\d*\.?\s*)?(.+)$", r"\1", subsection['heading'].strip())
                normalized = normalize_heading(clean_sub_title)
                subsection_map[normalized] = subsection

            # Füge Unterabschnitte in Schema-Reihenfolge hinzu
            for schema_subsection in schema_subsections:
                normalized_schema = normalize_heading(schema_subsection['title'])

                if normalized_schema in subsection_map:
                    subsection_data = subsection_map[normalized_schema]
                    clean_sub_title = re.sub(r"^(?:\d+\.?\d*\.?\s*)?(.+)$", r"\1", subsection_data['heading'].strip())
                    numbered_sub_title = f"{schema_subsection['id']} {clean_sub_title}"

                    result_lines.append(numbered_sub_title)
                    result_lines.extend(subsection_data['lines'])

    return "\n".join(result_lines)


def sanitize_sensitive_text(text):
    """Ersetzt Namen/Orte mit X. bzw. F., ohne [Anonymisiert] zu nutzen."""
    if not text:
        return text

    sanitized = text
    for pattern, repl in SENSITIVE_PATTERNS:
        sanitized = pattern.sub(repl, sanitized)
    return sanitized


def configure_document_styles(doc):
    """Konfiguriert die Dokumentstyles für einheitliche Formatierung.

    - Schriftgröße: 10pt für alle Styles
    - Heading 1: Abstand vor 12pt, nach 0pt
    - Heading 2: Abstand vor 12pt, nach 0pt
    - Heading 3: Abstand vor 6pt, nach 0pt
    - Normal: Abstand vor 6pt, nach 0pt
    """
    styles = doc.styles

    # Normal Style
    if 'Normal' in styles:
        normal_style = styles['Normal']
        normal_style.font.size = Pt(10)
        normal_style.paragraph_format.space_before = Pt(6)
        normal_style.paragraph_format.space_after = Pt(0)

    # Heading 1 Style
    if 'Heading 1' in styles:
        h1_style = styles['Heading 1']
        h1_style.font.size = Pt(10)
        h1_style.paragraph_format.space_before = Pt(12)
        h1_style.paragraph_format.space_after = Pt(0)

    # Heading 2 Style (Unterüberschriften wie 4.1, 6.1 - kein Abstand vor)
    if 'Heading 2' in styles:
        h2_style = styles['Heading 2']
        h2_style.font.size = Pt(10)
        h2_style.paragraph_format.space_before = Pt(0)
        h2_style.paragraph_format.space_after = Pt(0)

    # Heading 3 Style (Unterüberschriften - kein Abstand vor)
    if 'Heading 3' in styles:
        h3_style = styles['Heading 3']
        h3_style.font.size = Pt(10)
        h3_style.paragraph_format.space_before = Pt(0)
        h3_style.paragraph_format.space_after = Pt(0)

    # Heading 4 Style (falls vorhanden - kein Abstand vor)
    if 'Heading 4' in styles:
        h4_style = styles['Heading 4']
        h4_style.font.size = Pt(10)
        h4_style.paragraph_format.space_before = Pt(0)
        h4_style.paragraph_format.space_after = Pt(0)


def _insert_paragraph_after(paragraph, text, style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style:
        new_para.style = style
    return new_para


def detect_heading_level(text):
    """Erkennt die Überschriftenebene basierend auf Schema-Definitionen.

    Args:
        text: Der zu prüfende Text

    Returns:
        int oder None: 2 (h2/Heading 2), 3 (h3/Heading 3) oder None
    """
    cleaned = text.strip()

    # WICHTIG: Nur kurze Texte können Überschriften sein
    # Wenn der Text länger als 200 Zeichen ist, ist es sicher kein Heading
    if len(cleaned) > 200:
        return None

    # KRITISCH: Prüfe ZUERST auf Doppelpunkt - BEVOR Schema-Matching
    # Grund: "2.2 Ergebnisse...: Die BDI-II..." würde sonst als Heading erkannt
    # Heuristik: Wenn nach dem Doppelpunkt mehr als 40 Zeichen folgen, ist es Fließtext
    if ":" in cleaned:
        parts = cleaned.split(":", 1)
        if len(parts) == 2 and len(parts[1].strip()) > 40:
            # Das ist vermutlich "Überschrift: langer Fließtext"
            return None

    # Normalisiere für Vergleich
    normalized_input = normalize_heading(cleaned)

    # Exakte Übereinstimmung mit Schema (case-insensitive, whitespace-tolerant)
    for heading, level in HEADING_MAP.items():
        normalized_heading = normalize_heading(heading)

        # NUR exakter Match, NICHT "enthält"
        # Dies verhindert, dass "2.1 Auffälligkeiten: langer Text..." als Heading erkannt wird
        if normalized_input == normalized_heading:
            return level

    return None


def _replace_placeholder_with_text(doc, placeholder, text):
    """Ersetzt Platzhalter mit formatiertem Text.

    Überschriften werden automatisch mit passenden Styles versehen:
    - Hauptüberschriften: Heading 2
    - Unterüberschriften: Heading 3
    - Sub-Unterüberschriften: Heading 4
    - Normaler Text: Normal

    WICHTIG: Überschriften, die bereits im Template als Heading 1 existieren,
    werden NICHT nochmal eingefügt (verhindert Duplikate).
    """
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == placeholder:
            placeholder_para = paragraph  # Merke den ursprünglichen Platzhalter
            anchor = paragraph
            lines = [ln.rstrip() for ln in (text or "").splitlines()]
            while lines and not lines[0].strip():
                lines.pop(0)
            while lines and not lines[-1].strip():
                lines.pop()

            if not lines:
                _insert_paragraph_after(anchor, "[Angabe fehlt]")
            else:
                # Sammle alle existierenden Heading 1/2 Überschriften im Dokument
                existing_headings = set()
                for p in doc.paragraphs:
                    if p.style.name in ["Heading 1", "Heading 2"]:
                        normalized = normalize_heading(p.text)
                        existing_headings.add(normalized)

                for line in lines:
                    cleaned = line.strip()
                    if not cleaned:
                        anchor = _insert_paragraph_after(anchor, "")
                        continue

                    # Erkenne Überschriftenebene
                    heading_level = detect_heading_level(cleaned)

                    # Prüfe ob diese Überschrift bereits im Dokument existiert
                    if heading_level == 2:
                        normalized_current = normalize_heading(cleaned)
                        if normalized_current in existing_headings:
                            # Überspringe diese Zeile - Duplikat!
                            continue
                        style = "Heading 2"
                    elif heading_level == 3:
                        style = "Heading 3"
                    elif heading_level == 4:
                        style = "Heading 4"
                    else:
                        style = "Normal"

                    # KRITISCH: Aktualisiere anchor nach jedem Einfügen, um korrekte Reihenfolge zu garantieren
                    anchor = _insert_paragraph_after(anchor, cleaned, style=style)

            # Entferne den ursprünglichen Platzhalter (nicht den aktualisierten anchor!)
            placeholder_para._p.getparent().remove(placeholder_para._p)
            return


def _remove_leading_heading(text, heading):
    """Entfernt die führende Überschrift aus dem Text.

    Prüft normalisiert (ohne Nummerierung und case-insensitive), um
    Überschriften wie "1. Relevante soziodemographische Daten" zu erkennen,
    wenn heading="Relevante soziodemographische Daten".

    Behandelt zwei Fälle:
    1. Überschrift steht in separater Zeile: "1. Relevante...\nFrau X..."
    2. Überschrift am Anfang der ersten Zeile: "1. Relevante... Frau X..."
    """
    if not text or not heading:
        return text

    lines = [ln.rstrip() for ln in text.splitlines()]
    while lines and not lines[0].strip():
        lines.pop(0)

    if not lines:
        return text

    first_line = lines[0].strip()
    normalized_heading = normalize_heading(heading)

    # Fall 1: Erste Zeile ist EXAKT die Überschrift (normalisiert)
    normalized_first = normalize_heading(first_line)
    if normalized_first == normalized_heading:
        lines = lines[1:]
        # Entferne auch folgende Leerzeilen
        while lines and not lines[0].strip():
            lines.pop(0)
        return "\n".join(lines)

    # Fall 2: Erste Zeile BEGINNT mit der Überschrift
    # z.B. "Relevante soziodemographische Daten Frau X. ist..."
    # Suche nach der Überschrift am Anfang (mit/ohne Nummerierung)
    import re
    # Pattern: Optional "1." oder "1 " am Anfang, dann die Überschrift
    pattern = r'^(?:\d+\.?\s*)?' + re.escape(heading)
    match = re.match(pattern, first_line, re.IGNORECASE)
    if match:
        # Entferne den Match-Teil vom Anfang der ersten Zeile
        remaining = first_line[match.end():].strip()
        if remaining:
            # Es gibt noch Text nach der Überschrift in der gleichen Zeile
            lines[0] = remaining
        else:
            # Nur die Überschrift war in der Zeile
            lines = lines[1:]
            while lines and not lines[0].strip():
                lines.pop(0)

    return "\n".join(lines)


def ensure_report_template(template_path):
    """Erstellt das DOCX-Template, falls es nicht existiert."""
    if os.path.exists(template_path):
        return

    os.makedirs(os.path.dirname(template_path), exist_ok=True)
    doc = Document()

    doc.add_paragraph("Bericht an den Gutachter", style="Title")

    doc.add_paragraph("1. Relevante soziodemographische Daten", style="Heading 1")
    doc.add_paragraph("{{SECTION_1}}")

    doc.add_paragraph("2. Symptomatik und psychischer Befund", style="Heading 1")
    doc.add_paragraph("{{SECTION_2}}")

    doc.add_paragraph("3. Somatischer Befund / Konsiliarbericht", style="Heading 1")
    doc.add_paragraph("{{SECTION_3}}")

    doc.add_paragraph(
        "4. Behandlungsrelevante Angaben zur Lebensgeschichte, "
        "Krankheitsanamnese, funktionales Bedingungsmodell (VT)",
        style="Heading 1",
    )
    doc.add_paragraph("{{SECTION_4}}")

    doc.add_paragraph("5. Diagnose zum Zeitpunkt der Antragsstellung", style="Heading 1")
    doc.add_paragraph("{{SECTION_5}}")

    doc.add_paragraph("6. Behandlungsplan und Prognose", style="Heading 1")
    doc.add_paragraph("{{SECTION_6}}")

    # Konfiguriere Styles für einheitliche Formatierung
    configure_document_styles(doc)

    doc.save(template_path)


def create_comparison_docx(results, model_combinations, section_headers, parse_sections_func, enable_post_processing=True):
    """Erstellt eine DOCX-Datei mit Vergleichstabelle.

    Args:
        results: Liste mit Ergebnis-Strings (einer pro Modellkombination)
        model_combinations: Liste der Modellkombinationen [{"pass1": ..., "pass2": ...}, ...]
        section_headers: Liste der Abschnitts-Ueberschriften
        parse_sections_func: Funktion zum Parsen der Abschnitte aus dem Text
        enable_post_processing: Ob Post-Processing (Säuberung, Validierung, Repair) durchgeführt werden soll

    Returns:
        BytesIO mit dem DOCX-Dokument
    """
    doc = Document()

    num_combos = len(model_combinations)
    # Tabelle: Header + Abschnittszeilen, Spalten = 1 Ueberschrift + n Kombis + 1 Leer
    cols = num_combos + 2
    table = doc.add_table(rows=len(section_headers) + 1, cols=cols)
    table.style = "Table Grid"

    # Spaltenbreiten setzen (robust bei variabler Kombi-Anzahl)
    for row in table.rows:
        if cols > 0:
            row.cells[0].width = Cm(4)    # Ueberschriften
        for idx in range(1, cols - 1):
            row.cells[idx].width = Cm(4)  # Kombis
        row.cells[cols - 1].width = Cm(1)  # Leer

    # Header-Zeile (erste Zeile leer in Spalte 1)
    header_row = table.rows[0]
    header_row.cells[0].text = ""
    for idx, combo in enumerate(model_combinations, start=1):
        header_row.cells[idx].text = f"{combo['pass1']} + {combo['pass2']}"
    header_row.cells[cols - 1].text = ""

    # Extrahiere Abschnitte aus jedem Ergebnis
    parsed_results = [parse_sections_func(result) for result in results]

    # Fuelle die Tabelle (Zeilen 2..n)
    for row_idx in range(1, len(section_headers) + 1):
        section_idx = row_idx - 1
        row = table.rows[row_idx]

        # Spalte 1: Ueberschrift
        row.cells[0].text = section_headers[section_idx]

        # Spalten 2..(n+1): Ergebnisse (eine pro Kombination)
        for combo_idx in range(num_combos):
            cell_text = parsed_results[combo_idx][section_idx]

            # Post-Processing-Pipeline anwenden (falls aktiviert)
            if enable_post_processing:
                pp_result = post_process_text(
                    cell_text,
                    enable_repair=True,
                    enable_validation=False
                )
                cell_text = pp_result["text"]
            else:
                # Legacy-Pfad ohne Post-Processing
                cell_text = sanitize_sensitive_text(cell_text)

            row.cells[combo_idx + 1].text = cell_text

        # Letzte Spalte: leer
        row.cells[cols - 1].text = ""

    # Formatierung: 10pt Schriftgröße, keine Abstände in Tabellenzellen
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    # In BytesIO speichern
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output


def format_text_as_html(text):
    """Formatiert Text mit HTML-Tags für Überschriften (h2-h4).

    Args:
        text: Der zu formatierende Text

    Returns:
        str: HTML-formatierter Text
    """
    if not text:
        return ""

    lines = text.splitlines()
    html_lines = []

    for line in lines:
        cleaned = line.strip()
        if not cleaned:
            html_lines.append("<p></p>")
            continue

        # Erkenne Überschriftenebene
        heading_level = detect_heading_level(cleaned)
        if heading_level == 2:
            html_lines.append(f"<h2>{cleaned}</h2>")
        elif heading_level == 3:
            html_lines.append(f"<h3>{cleaned}</h3>")
        elif heading_level == 4:
            html_lines.append(f"<h4>{cleaned}</h4>")
        else:
            html_lines.append(f"<p>{cleaned}</p>")

    return "\n".join(html_lines)


def create_flowing_text_docx(sections, selected_texts, enable_post_processing=True):
    """Erstellt eine DOCX-Datei mit Fliesstext aus ausgewaehlten Abschnitten.

    Args:
        sections: Liste der Abschnitts-Ueberschriften
        selected_texts: Liste der ausgewaehlten Texte (einer pro Abschnitt)
        enable_post_processing: Ob Post-Processing (Säuberung, Validierung, Repair) durchgeführt werden soll

    Returns:
        BytesIO mit dem DOCX-Dokument
    """
    template_path = os.path.join(os.path.dirname(__file__), "templates", TEMPLATE_FILENAME)
    ensure_report_template(template_path)
    doc = Document(template_path)

    # Konfiguriere Styles für einheitliche Formatierung (10pt, korrekte Abstände)
    configure_document_styles(doc)

    for idx, section_text in enumerate(selected_texts, 1):
        placeholder = f"{{{{SECTION_{idx}}}}}"
        heading = sections[idx - 1] if idx - 1 < len(sections) else None

        # Post-Processing-Pipeline anwenden (falls aktiviert)
        if enable_post_processing:
            # WICHTIG: enable_repair=False, weil repair_schema() erwartet den KOMPLETTEN Bericht
            # mit allen 6 Abschnitten, aber wir haben hier nur EINEN Abschnitt!
            pp_result = post_process_text(
                section_text,
                enable_repair=False,  # Repair nur für Gesamtdokument, nicht einzelne Abschnitte
                enable_validation=False
            )
            section_text = pp_result["text"]

            # Log repair actions (falls vorhanden)
            if pp_result["repair_log"]:
                print(f"[Post-Processing] Abschnitt {idx}: {len(pp_result['repair_log'])} Reparaturen")
                for log_entry in pp_result["repair_log"]:
                    print(f"  - {log_entry}")
        else:
            # Legacy-Pfad ohne Post-Processing
            section_text = sanitize_sensitive_text(section_text)

        section_text = _remove_leading_heading(section_text, heading)
        _replace_placeholder_with_text(doc, placeholder, section_text)

    # In BytesIO speichern
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output
